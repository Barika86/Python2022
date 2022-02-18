# -*- coding: utf-8 -*-
"""
Created on Fri Jan  7 21:43:41 2022

@author: Ragab Barika
"""

from ast import Break
from turtle import width
from docx import Document # to import documents
import pyttsx3 # for speak 


def speak(text):
    pyttsx3.speak(text)


Document = Document()

# resize a picture 
from docx.shared import Inches 

# adding a profile picture 
Document.add_picture('me.JPEG', width=Inches(2.0))

# name, phone number and details
name = input('what is your name? ')
speak('hello ' + name + 'how are you today?')

speak('what is your phone number?')
phone_number = input('what is your phone_number? ')
speak('what is your email?')
email = input('what is your email ')

Document.add_paragraph(name + '|' + phone_number + '|' + email)





# about me 
# Headings and about section
Document.add_heading('about me')
speak('tell us about yourself?' )
Document.add_paragraph(input('tell us about yourself?' ))

# Work experiences
Document.add_heading('Work Experience')
p = Document.add_paragraph()
company = input ('Enter Company ')
from_date = input ('From Date ')
to_date = input ('To Date ')

p.add_run(company + ' ' ).blod = True # to add text to exist to pragraph
p.add_run(from_date + '-' + to_date + '\n').italic = True 

experienc_details = input('describe your experience at ' + company )
p.add_run (experienc_details)

# More experiences 

while True:
    has_more_experiences = input ('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes': 
        p = Document.add_paragraph()

        company = input ('Enter Company ')
        from_date = input ('From Date ')
        to_date = input ('To date ')

        p.add_run(company + ' ' ).bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True 

        experienc_details = input('describe your experience at ' + company + ' ' )
        p.add_run (experienc_details)
    else:
        break

#More skills 
Document.add_heading('Skills')
Skill = input ('Enter Skill ')
p = Document.add_paragraph(Skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more Skills ? Yes or No ')
    if has_more_skills.lower() == 'yes':
        Skill = input ('Enter skill ') 
        p = Document.add_paragraph(Skill)
        p.style = 'List Bullet'
    else:
        break



# footer
section = Document.sections[0]
footer =section.footer
p = footer.paragraphs[0]
p.text =' CV generated using Ragabs laptop'

Document.save('cv.docx')

