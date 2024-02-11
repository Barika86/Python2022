
#  Create Word Documents 

#  Write Paragraphs to Word Documents with Python

from docx import Document
from docx.shared import Inches
import pyttsx3  ### Python Text to Speech 

def speak (text):
    pyttsx3.speak(text) 

# Create a new Document object
document = Document()

# Insert profile picture
document.add_picture('me.jpeg', width=Inches(2.0))

# Define personal information variables
name = input('What is your name? ')
speak ('Hello' + name + 'how are you today ') ## Python Text to Speech 


speak ( 'What is your phone number ')
Phone_number = input('What is your phone number? ')


speak ('What is your email address ')
Email = input('What is your email address? ')




# Add personal information to the document
document.add_paragraph(name + ' | ' + Phone_number + ' | ' + Email) 

# Write about me section
document.add_heading('About me', level=1)

speak ('Tell me about yourself ')
about_me_text = input('Tell me about yourself: ')
document.add_paragraph(about_me_text)

# Write Work Experience section
speak ('Tell me about your Experience ')
document.add_heading('Work Experience', level=2)

# Input and add details for the first work experience
company = input('Enter company: ')
from_date = input('From Date: ')
to_date = input('To Date: ')

p = document.add_paragraph()
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Describe your experience at ' + company + ': ')
p.add_run(experience_details)

# Loop for adding more work experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No: ')
    if has_more_experiences.lower() == 'yes':
        company = input('Enter the company name: ')
        from_date = input('Enter the start date: ')
        to_date = input('Enter the end date: ')

        p = document.add_paragraph()
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Describe your experience at ' + company + ': ')
        p.add_run(experience_details)
    else:
        break

# Skills Section
document.add_heading('Skills')
skills = input('Enter skills: ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter the name of skills: ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

# Footer 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text  = 'CV generated using Ragab computer'

document.save('cv.docx')
